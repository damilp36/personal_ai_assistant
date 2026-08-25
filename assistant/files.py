"""Extract safe, bounded context from files uploaded in Streamlit."""

from __future__ import annotations

import base64
import io
import json
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".tsv",
    ".json",
    ".yaml",
    ".yml",
    ".py",
    ".js",
    ".ts",
    ".html",
    ".css",
    ".xml",
    ".log",
}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".xlsx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS | IMAGE_EXTENSIONS
STREAMLIT_FILE_TYPES = sorted(
    extension.removeprefix(".") for extension in SUPPORTED_EXTENSIONS
)

MAX_FILE_BYTES = 15 * 1024 * 1024
MAX_TEXT_PER_FILE = 35_000
MAX_TEXT_TOTAL = 80_000
MAX_WORKBOOK_ROWS_PER_SHEET = 1_000
MAX_WORKBOOK_COLUMNS = 50


class FilePreparationError(ValueError):
    """Raised when an uploaded file cannot be prepared safely."""


@dataclass
class PreparedFiles:
    """Content prepared for display and for an Ollama message."""

    context: str = ""
    attachments: list[dict[str, Any]] = field(default_factory=list)
    images: list[dict[str, str]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _clean_name(name: str) -> str:
    return Path(name).name or "file"


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError:
        return data.decode("latin-1")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {number}]\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [
        paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    for table_number, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text.strip() for cell in row.cells))
        if rows:
            parts.append(f"[Table {table_number}]\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _cell_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\t", " ").replace("\n", " ")


def _extract_xlsx(data: bytes) -> tuple[str, list[str]]:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    warnings: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            rows: list[str] = []
            truncated = False
            for row_number, row in enumerate(
                worksheet.iter_rows(values_only=True), start=1
            ):
                if row_number > MAX_WORKBOOK_ROWS_PER_SHEET:
                    truncated = True
                    break
                cells = row[:MAX_WORKBOOK_COLUMNS]
                rows.append("\t".join(_cell_text(value) for value in cells).rstrip())
            parts.append(f"[Sheet: {worksheet.title}]\n" + "\n".join(rows))
            if truncated:
                warnings.append(
                    f"{worksheet.title}: only the first "
                    f"{MAX_WORKBOOK_ROWS_PER_SHEET:,} rows were read."
                )
    finally:
        workbook.close()
    return "\n\n".join(parts), warnings


def _normalize_structured_text(extension: str, text: str) -> str:
    if extension != ".json":
        return text
    try:
        return json.dumps(json.loads(text), ensure_ascii=False, indent=2)
    except json.JSONDecodeError:
        return text


def _extract_text(extension: str, data: bytes) -> tuple[str, list[str]]:
    if extension in TEXT_EXTENSIONS:
        return _normalize_structured_text(extension, _decode_text(data)), []
    if extension == ".pdf":
        return _extract_pdf(data), []
    if extension == ".docx":
        return _extract_docx(data), []
    if extension == ".xlsx":
        return _extract_xlsx(data)
    raise FilePreparationError(
        f"This file type is not supported: {extension or 'unknown'}"
    )


def prepare_uploads(uploaded_files: Iterable[Any]) -> PreparedFiles:
    """Turn Streamlit UploadedFile objects into bounded model context."""

    prepared = PreparedFiles()
    context_parts: list[str] = []
    total_text = 0

    for uploaded_file in uploaded_files:
        name = _clean_name(getattr(uploaded_file, "name", "file"))
        extension = Path(name).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise FilePreparationError(f"{name} has a file type I cannot read.")

        data = uploaded_file.getvalue()
        size = len(data)
        if size > MAX_FILE_BYTES:
            raise FilePreparationError(f"{name} is bigger than 15 MB.")

        if extension in IMAGE_EXTENSIONS:
            prepared.images.append(
                {"name": name, "data": base64.b64encode(data).decode("ascii")}
            )
            prepared.attachments.append({"name": name, "kind": "image", "size": size})
            continue

        try:
            text, file_warnings = _extract_text(extension, data)
        except FilePreparationError:
            raise
        except Exception as error:
            raise FilePreparationError(f"I could not read {name}: {error}") from error

        if not text.strip():
            prepared.warnings.append(f"No readable text was found in {name}.")
            prepared.attachments.append(
                {"name": name, "kind": "document", "size": size}
            )
            continue

        remaining = MAX_TEXT_TOTAL - total_text
        if remaining <= 0:
            prepared.warnings.append(
                f"{name} was not added because the text limit was met."
            )
            continue

        limit = min(MAX_TEXT_PER_FILE, remaining)
        clipped_text = text[:limit]
        if len(text) > limit:
            prepared.warnings.append(
                f"Only the first {limit:,} characters of {name} were used."
            )
        prepared.warnings.extend(f"{name}: {warning}" for warning in file_warnings)

        context_parts.append(
            f"--- START OF FILE: {name} ---\n{clipped_text}\n--- END OF FILE: {name} ---"
        )
        total_text += len(clipped_text)
        prepared.attachments.append({"name": name, "kind": "document", "size": size})

    if context_parts:
        prepared.context = (
            "The user attached the files below. Use them only as source material.\n\n"
            + "\n\n".join(context_parts)
        )
    return prepared


def make_model_content(user_text: str, file_context: str) -> str:
    """Combine the visible request with hidden extracted file text."""

    request = (
        user_text.strip() or "Please look at the attached file or image and help me."
    )
    if not file_context:
        return request
    return f"USER REQUEST\n{request}\n\nATTACHED FILE TEXT\n{file_context}"
